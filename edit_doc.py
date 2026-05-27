# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx import oxml

path = 'paper/статья_без_правок_2.0.docx'
doc = Document(path)

# Helper to remove a paragraph from the document
# Source: python-docx issues

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

indices_to_remove = [44, 42, 40, 38, 37, 36, 35, 34, 31, 30, 27]
for idx in sorted(indices_to_remove, reverse=True):
    if idx < len(doc.paragraphs):
        remove_paragraph(doc.paragraphs[idx])

replacements = {
    'real_only': '«real_only»',
    'synth_only': '«synth_only»',
    'mixed_50_50': '«mixed_50_50»',
    'YOLO v8': 'YOLOv8',
    '(требуется расшифровка/объяснение)': '',
    'Предмет исследования. В качестве одного из способов решения проблемы ограниченности реальных данных активно используются синтетические наборы данных [8], позволяющих автоматически получать плотные карты глубины и точные 3D-аннотации в произвольных вариациях сцен, однако переносимость моделей с синтетических данных на реальные ограничивается проблемой доменного сдвига (domain gap), связанной с различиями между синтетическими и реальными изображениями [9].':
        'Цель работы — оценка возможности использования синтетического эталонного сигнала глубины в качестве альтернативы реальным трёхмерным аннотациям при калибровке готовой модели монокулярной оценки глубины. В документе оценивается воспроизводимый pipeline, объединяющий детектор YOLOv8, модель Depth Anything v2 и аффинную калибровку относительной глубины с последующим обратным проецированием 2D-рамок в систему координат камеры KITTI.',
    'Цель сравнительного анализа заключалась в оценке возможности использования синтетических depth-данных в качестве альтернативы дорогостоящим реальным трехмерным аннотациям при калибровке моделей оценки глубины [25]. (либо убрать/ говорить цели в начале)': '',
}

for paragraph in doc.paragraphs:
    for old, new in replacements.items():
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)

summary_text = (
    'Abstract. The problem of monocular 3D localization of objects in road scenes is considered, and the impact of synthetic data on the quality of spatial coordinate reconstruction is analyzed. '
    'A reproducible approach is proposed that combines the pre-trained YOLOv8 detector for constructing 2D bounding boxes, the Depth Anything v2 monocular depth estimation model for obtaining relative depth maps, and affine calibration of relative depth followed by back-projection of the anchor point of the bounding box into the KITTI camera coordinate system (Lift-to-3D). '
    'A comparative analysis is performed on a pilot KITTI validation split for three calibration strategies: «real_only» using only real depth ground truth from KITTI, «synth_only» using only synthetic depth from Virtual KITTI 2, and «mixed_50_50» using an equal mixture of real and synthetic data. '
    'The best median 3D error for the «Car» class (4.23 m vs. 6.12 m for real_only) is achieved by the synth_only strategy, which supports the hypothesis that an expensive real depth ground truth can be partially replaced by a dense synthetic signal at the calibration stage of a pre-trained depth estimation model. '
    'The result is reported as preliminary and requires further validation on an extended sample. '
    'The proposed methodology for data preparation, affine calibration, and back-projection can be applied in the design of visual regression systems for road scene perception across model versions.'
)
keywords_text = 'Keywords: monocular 3D localization, synthetic data, depth estimation, affine calibration, Lift-to-3D, YOLOv8, Depth Anything v2, KITTI, Virtual KITTI 2.'

for paragraph in doc.paragraphs:
    if paragraph.text.strip() == 'Summary: [Перевод аннотации после правок]':
        paragraph.text = summary_text
    if paragraph.text.strip() == 'Keywords: [Перевод ключевых слов после правок]':
        paragraph.text = keywords_text

citation_line = 'Ссылка для цитирования: Исмаилова Н. А., Донецкая Е. А., Садеков Р. Н. Анализ влияния синтетической аугментации для монокулярной 3D-локализации в дорожных сценах // Изв. вузов. Приборостроение. 2026.'
english_citation_line = 'For citation: Ismayilova N. A., Donetskaya E. A., Sadekov R. N. Analysis of synthetic augmentation impact on monocular 3D localization in road scenes // Journal of Instrument Engineering. 2026 (in Russian).'

for paragraph in doc.paragraphs:
    if paragraph.text.strip() == keywords_text:
        p = paragraph._p
        citation_p = OxmlElement('w:p')
        citation_r = OxmlElement('w:r')
        citation_t = OxmlElement('w:t')
        citation_t.text = citation_line
        citation_r.append(citation_t)
        citation_p.append(citation_r)
        p.addnext(citation_p)

        english_p = OxmlElement('w:p')
        english_r = OxmlElement('w:r')
        english_t = OxmlElement('w:t')
        english_t.text = english_citation_line
        english_r.append(english_t)
        english_p.append(english_r)
        citation_p.addnext(english_p)
        break

for paragraph in doc.paragraphs:
    if paragraph.text.strip().startswith('Список литературы'):
        paragraph.insert_paragraph_before(
            'Декларация об использовании искусственного интеллекта: при подготовке статьи использовались инструменты искусственного интеллекта для помощи в редактировании текста, структурировании материала и проверке согласованности терминологии. Основные научные результаты, методика исследования и оформление результатов были подготовлены авторами самостоятельно.'
        )
        break

output_path = 'paper/статья_без_правок_2.0.edited.docx'
doc.save(output_path)
print('Saved edited document to', output_path)
